"""
其他格式解析器（Markdown / TXT / Word / 网页）

作用：
    提供除 PDF 外的其他文档格式解析能力。
    PDF 解析逻辑复杂，单独放在 pdf_parser.py 中。

    各解析器统一返回 raw_text，由后续步骤（清洗、分块）统一处理。

实现方式：
    1. MarkdownParser - 直接读取（Markdown 本质是纯文本）
    2. TxtParser - 多编码尝试（UTF-8 → GBK → Latin-1）
    3. DocxParser - python-docx 提取段落和表格
    4. UrlParser - requests + BeautifulSoup 提取正文
"""

import os
import logging
from typing import List, Optional

import requests
from bs4 import BeautifulSoup
from docx import Document as DocxDocument

from app.services.document_pipeline.context import (
    PipelineContext,
    ParsedPage,
    TextBlock,
)

logger = logging.getLogger(__name__)


# ============================================
# Markdown 解析器
# ============================================

class MarkdownParser:
    """
    Markdown 解析器

    作用：
        读取 Markdown 文件的原始内容。
        Markdown 本质是纯文本，无需特殊解析，直接读取即可。

    使用方式：
        parser = MarkdownParser()
        text = parser.parse("/path/to/doc.md")
    """

    def parse(self, file_path: str) -> str:
        """
        解析 Markdown 文件

        作用：
            以 UTF-8 编码读取 Markdown 文件。

        参数：
            file_path: str - 文件路径

        返回：
            str - 文件内容
        """
        with open(file_path, "r", encoding="utf-8") as f:
            return f.read()

    def parse_to_context(self, ctx: PipelineContext) -> None:
        """
        解析并填充上下文

        作用：
            读取文件内容写入 ctx.raw_text，并把整篇文档作为单页放入 ctx.pages。
            Markdown 没有分页概念，作为单页处理。

        参数：
            ctx: PipelineContext - 流水线上下文
        """
        text = self.parse(ctx.file_path)
        ctx.raw_text = text
        ctx.pages = [
            ParsedPage(
                page_number=1,
                text=text,
                blocks=[TextBlock(text=text, page_number=1)],
                is_blank=not text.strip(),
            )
        ]


# ============================================
# TXT 解析器
# ============================================

class TxtParser:
    """
    纯文本解析器

    作用：
        读取 TXT 文件，自动检测编码。
        Windows 中文环境常见 GBK 编码，需多编码尝试。

    使用方式：
        parser = TxtParser()
        text = parser.parse("/path/to/doc.txt")
    """

    # 尝试的编码列表（按优先级）
    # 作用：UTF-8 是首选，失败则尝试 GBK（Windows 中文），最后用 Latin-1 兜底
    _ENCODINGS = ["utf-8", "utf-8-sig", "gbk", "gb2312", "big5", "latin-1"]

    def parse(self, file_path: str) -> str:
        """
        解析 TXT 文件

        作用：
            按优先级尝试多种编码，找到第一个能正确解码的编码。

        实现方式：
            1. 依次尝试 _ENCODINGS 中的编码
            2. 第一个不抛 UnicodeDecodeError 的编码即为正确编码
            3. 全部失败则用 Latin-1 强制解码（不抛异常但可能有乱码）

        参数：
            file_path: str - 文件路径

        返回：
            str - 文件内容
        """
        with open(file_path, "rb") as f:
            raw_bytes = f.read()

        for encoding in self._ENCODINGS:
            try:
                return raw_bytes.decode(encoding)
            except UnicodeDecodeError:
                continue

        # 全部失败，用 Latin-1 强制解码（不会抛异常）
        logger.warning(f"文件 {file_path} 编码检测失败，使用 latin-1 兜底")
        return raw_bytes.decode("latin-1", errors="replace")

    def parse_to_context(self, ctx: PipelineContext) -> None:
        """
        解析并填充上下文

        作用：
            读取文件内容写入 ctx.raw_text 和 ctx.pages。

        参数：
            ctx: PipelineContext - 流水线上下文
        """
        text = self.parse(ctx.file_path)
        ctx.raw_text = text
        ctx.pages = [
            ParsedPage(
                page_number=1,
                text=text,
                blocks=[TextBlock(text=text, page_number=1)],
                is_blank=not text.strip(),
            )
        ]


# ============================================
# Word（DOCX）解析器
# ============================================

class DocxParser:
    """
    Word 文档解析器

    作用：
        使用 python-docx 提取 Word 文档的段落和表格。
        DOCX 是 OOXML 格式，python-docx 能稳定解析。

    使用方式：
        parser = DocxParser()
        text = parser.parse("/path/to/doc.docx")
    """

    def parse(self, file_path: str) -> str:
        """
        解析 Word 文档

        作用：
            提取文档中所有段落和表格文本，按顺序拼接。

        实现方式：
            1. DocxDocument 打开文档
            2. 遍历 document.element.body 按原顺序处理段落和表格
            3. 表格转为 Markdown 格式插入文本流

        参数：
            file_path: str - 文件路径

        返回：
            str - 提取的文本
        """
        doc = DocxDocument(file_path)
        parts: List[str] = []

        # 按文档原始顺序遍历段落和表格
        # 作用：python-docx 的 doc.paragraphs 和 doc.tables 是分开的列表，
        # 需要通过底层元素遍历才能保持原顺序
        from docx.oxml.ns import qn
        body = doc.element.body
        for child in body.iterchildren():
            if child.tag == qn("w:p"):
                # 段落
                from docx.text.paragraph import Paragraph
                para = Paragraph(child, doc)
                if para.text.strip():
                    parts.append(para.text)
            elif child.tag == qn("w:tbl"):
                # 表格
                from docx.table import Table
                table = Table(child, doc)
                parts.append(self._table_to_markdown(table))

        return "\n\n".join(parts)

    def _table_to_markdown(self, table) -> str:
        """
        将 Word 表格转为 Markdown 格式

        作用：
            把 docx 表格转为 Markdown 表格字符串，便于后续统一处理。

        参数：
            table: docx.table.Table - docx 表格对象

        返回：
            str - Markdown 表格字符串
        """
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append("| " + " | ".join(cells) + " |")

        if len(rows) >= 1:
            # 在第一行后插入分隔行
            separator = "| " + " | ".join(["---"] * len(table.rows[0].cells)) + " |"
            rows.insert(1, separator)

        return "\n".join(rows)

    def parse_to_context(self, ctx: PipelineContext) -> None:
        """
        解析并填充上下文

        作用：
            提取文本写入 ctx.raw_text，作为单页处理。

        参数：
            ctx: PipelineContext - 流水线上下文
        """
        text = self.parse(ctx.file_path)
        ctx.raw_text = text
        ctx.pages = [
            ParsedPage(
                page_number=1,
                text=text,
                blocks=[TextBlock(text=text, page_number=1)],
                is_blank=not text.strip(),
            )
        ]


# ============================================
# 网页解析器
# ============================================

class UrlParser:
    """
    网页解析器

    作用：
        下载网页并提取正文内容，去除导航、广告等噪声。

    使用方式：
        parser = UrlParser()
        text = parser.parse("https://example.com/article")
    """

    # 请求头，模拟真实浏览器访问
    # 作用：部分网站（如百度百科）反爬机制会拒绝请求头不完整的请求，
    #       返回 403/412 等状态码。完整请求头可提高网页导入成功率。
    _HEADERS = {
        "User-Agent": (
            "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
            "AppleWebKit/537.36 (KHTML, like Gecko) "
            "Chrome/120.0.0.0 Safari/537.36"
        ),
        "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8",
        "Accept-Language": "zh-CN,zh;q=0.9,en;q=0.8",
        "Accept-Encoding": "gzip, deflate",
        "Connection": "keep-alive",
        "Upgrade-Insecure-Requests": "1",
    }

    def _download_html(self, url: str, timeout: int = 30, _redirect_count: int = 0) -> tuple:
        """
        下载网页 HTML（内部方法，供 parse 和 parse_with_images 复用）

        作用：
            执行带 SSRF 防护的 HTTP 下载，返回响应内容和 response 对象。
            调用方负责关闭 response 和检测 Content-Type。

        安全防护：
            - C-1: 禁用重定向（防止 SSRF 绕过）
            - C-8: 流式下载 + 大小限制（防止 OOM）

        参数：
            url: str - 网页 URL（已经过 validate_url 校验）
            timeout: int - 超时时间（秒）

        返回:
            tuple[bytes, requests.Response] - (响应内容, response 对象)
            调用方负责调用 response.close()

        异常：
            ValueError - 重定向或大小超限时抛出
            requests.RequestException - 网络/HTTP 错误
        """
        # C-1 + C-8 修复：禁用自动重定向 + 流式下载 + 大小限制
        # 安全策略：不使用 requests 自动重定向，手动跟随并校验每个重定向目标
        response = requests.get(
            url,
            headers=self._HEADERS,
            timeout=timeout,
            allow_redirects=False,  # C-1: 禁用自动重定向，手动校验后跟随
            stream=True,            # C-8: 流式下载，便于边读边检查大小
        )

        # C-1: 安全的重定向跟随（手动校验每个目标 URL）
        # 作用：部分网站（如百度百科）会 302 重定向到规范 URL，
        #       完全禁止重定向会导致导入失败；但自动重定向有 SSRF 风险
        #       （重定向到内网地址），折中方案：手动获取 Location →
        #       SSRF 校验 → 递归调用，最多 3 次
        #       未通过 SSRF 校验的重定向目标：禁止 URL 重定向到不安全地址
        if response.is_redirect or response.is_permanent_redirect:
            response.close()

            if _redirect_count >= 3:
                raise ValueError(
                    f"URL 重定向次数超过限制（最多 3 次，target={url[:100]}）"
                )

            location = response.headers.get("Location")
            if not location:
                raise ValueError("重定向响应缺少 Location 头")

            # 解析相对路径为绝对 URL
            from urllib.parse import urljoin
            redirect_url = urljoin(url, location)

            # SSRF 校验重定向目标（与初始 URL 同等防护）
            from app.core.url_validator import validate_url, URLValidationError
            from app.core.config import settings as _settings
            try:
                validate_url(redirect_url, allow_private=_settings.is_development)
            except URLValidationError:
                raise ValueError(
                    f"重定向目标不安全，已拦截: {redirect_url[:100]}"
                )

            logger.info(
                f"URL 重定向跟随（{_redirect_count + 1}/3）: "
                f"{url[:80]} → {redirect_url[:80]}"
            )
            return self._download_html(redirect_url, timeout, _redirect_count + 1)

        response.raise_for_status()

        # C-8: 检查 Content-Length（若服务端提供）
        from app.core.config import settings
        max_size = settings.URL_IMPORT_MAX_SIZE
        content_length = response.headers.get("Content-Length")
        if content_length and int(content_length) > max_size:
            response.close()
            raise ValueError(
                f"下载内容过大（{int(content_length)} bytes），最大允许 {max_size} bytes"
            )

        # C-8: 流式读取 + 边写边检查累计大小
        downloaded = 0
        chunks = []
        for chunk in response.iter_content(chunk_size=1024 * 1024):  # 1MB chunks
            if not chunk:
                continue
            downloaded += len(chunk)
            if downloaded > max_size:
                response.close()
                raise ValueError(
                    f"下载内容超过大小限制（{max_size} bytes），已中止下载"
                )
            chunks.append(chunk)

        content = b"".join(chunks)
        # 注意：不在此处 close response，调用方需要读取 headers（如 Content-Type）
        return content, response

    @staticmethod
    def _is_html_content(content_type: str) -> bool:
        """
        判断响应是否为 HTML 内容

        作用：
            检测 Content-Type 是否为 text/html 或 application/xhtml+xml。
            用于拦截图片、PDF、视频等非网页资源的导入。
            无 Content-Type 时默认按 HTML 处理（兼容部分服务器不返回该头）。

        参数：
            content_type: str - HTTP 响应的 Content-Type 头（已转小写）

        返回:
            bool - True 表示是 HTML 内容
        """
        if not content_type:
            return True
        return (
            "text/html" in content_type
            or "application/xhtml+xml" in content_type
            or "text/plain" in content_type  # 部分服务器对 HTML 返回 text/plain
        )

    def parse(self, url: str, timeout: int = 30) -> str:
        """
        下载并解析网页（纯文本提取）

        作用：
            下载网页 HTML，提取正文文本。
            含 Content-Type 检测：非 HTML 资源抛出 ValueError("URL_NOT_HTML:...")。

        实现方式：
            1. _download_html 下载页面（含 SSRF 防护）
            2. 检测 Content-Type，非 HTML 抛出明确异常
            3. BeautifulSoup 解析 HTML
            4. 移除 script/style/nav 等非正文标签
            5. 提取纯文本，去除多余空白

        参数：
            url: str - 网页 URL
            timeout: int - 超时时间（秒）

        返回:
            str - 提取的正文文本

        异常:
            ValueError - 重定向/大小超限/非 HTML 资源时抛出
                非 HTML 资源异常消息含 "URL_NOT_HTML" 标识
            requests.RequestException - 下载失败
        """
        content, response = self._download_html(url, timeout)

        # 任务3：Content-Type 检测
        # 作用：URL 指向图片等非 HTML 资源时，BeautifulSoup 解析为空文本
        #       会被误判为"网页内容为空"→ HTTP 400，错误信息不友好
        #       修复：检测 Content-Type，非 HTML 时抛出带 URL_NOT_HTML 标识的异常
        content_type = response.headers.get("Content-Type", "").lower()
        response.close()

        if not self._is_html_content(content_type):
            raise ValueError(
                f"URL_NOT_HTML:URL 指向非 HTML 资源（Content-Type: {content_type}），"
                f"无法作为文档导入"
            )

        # BeautifulSoup 解析
        soup = BeautifulSoup(content, "html.parser")

        # 移除非正文标签
        for tag in soup(["script", "style", "nav", "footer", "header", "aside"]):
            tag.decompose()

        # 提取文本
        text = soup.get_text(separator="\n", strip=True)

        # 去除多余空行
        lines = [line.strip() for line in text.splitlines() if line.strip()]
        return "\n".join(lines)

    def parse_with_images(self, url: str, timeout: int = 30) -> str:
        """
        下载网页并提取正文 + 图片文本（任务3）

        作用：
            一次下载网页，提取纯文本和图片文本，拼接返回。
            避免重复下载。图片文本由 OCR/多模态模型生成。

        实现方式:
            1. _download_html 下载页面（含 SSRF 防护）
            2. 检测 Content-Type，非 HTML 抛出 URL_NOT_HTML 异常
            3. BeautifulSoup 解析 HTML
            4. 移除非正文标签
            5. 提取正文文本
            6. 提取图片文本（OCR + 多模态描述）
            7. 拼接正文 + 图片文本返回

        参数：
            url: str - 网页 URL
            timeout: int - 超时时间（秒）

        返回:
            str - 正文文本 + 图片文本

        异常:
            ValueError - 重定向/大小超限/非 HTML 资源时抛出
            requests.RequestException - 下载失败
        """
        content, response = self._download_html(url, timeout)

        # Content-Type 检测
        content_type = response.headers.get("Content-Type", "").lower()
        response.close()

        if not self._is_html_content(content_type):
            raise ValueError(
                f"URL_NOT_HTML:URL 指向非 HTML 资源（Content-Type: {content_type}）"
            )

        soup = BeautifulSoup(content, "html.parser")

        # 移除非正文标签
        for tag in soup(["script", "style", "nav", "footer", "header", "aside"]):
            tag.decompose()

        # 提取正文文本
        text = soup.get_text(separator="\n", strip=True)
        lines = [line.strip() for line in text.splitlines() if line.strip()]
        text = "\n".join(lines)

        # 提取图片文本
        # 作用：将网页中的图片转为文本（OCR/描述），纳入检索范围
        # 失败不影响文本导入（内部 catch all）
        try:
            from app.services.document_pipeline.image_processor import ImageProcessor
            image_processor = ImageProcessor()
            image_text = self.extract_image_texts(soup, url, image_processor)
            if image_text:
                separator = "=" * 40
                text = (
                    text + "\n\n" + separator + "\n图片内容\n" + separator + "\n" + image_text
                )
        except Exception as e:
            logger.warning(f"网页图片提取失败，跳过图片文本: {e}")

        return text

    # ============================================
    # 网页图片提取（任务3）
    # ============================================

    def extract_image_texts(
        self,
        soup: BeautifulSoup,
        base_url: str,
        image_processor: Optional[object] = None,
    ) -> str:
        """
        从网页 HTML 中提取图片并转为文本

        作用：
            遍历 <img> 标签，下载图片，过滤小图片，
            使用 OCR/多模态模型生成文本描述，附加到正文。
            单张图片失败不影响整体导入。

        实现方式：
            1. 解析 <img> 标签，提取 src 属性
            2. 相对 URL 转绝对 URL
            3. SSRF 校验（复用 url_validator）
            4. 逐张下载，检测尺寸，过滤小图片
            5. 调用 ImageProcessor 生成 OCR/描述文本
            6. 拼接所有图片文本返回

        参数：
            soup: BeautifulSoup - 已解析的 HTML
            base_url: str - 网页基础 URL（用于解析相对路径）
            image_processor: Optional[ImageProcessor] - 图片处理器
                为 None 则跳过 OCR/Vision，仅返回图片 URL 元信息

        返回:
            str - 拼接的图片文本（无图片时返回空字符串）
        """
        from app.core.config import settings
        from urllib.parse import urljoin

        img_tags = soup.find_all("img")
        if not img_tags:
            return ""

        # 限制图片数量，避免过多图片拖慢导入
        img_tags = img_tags[: settings.URL_IMAGE_MAX_COUNT]

        image_texts: List[str] = []
        for idx, img in enumerate(img_tags, start=1):
            src = img.get("src") or img.get("data-src")
            if not src:
                continue

            # 跳过 data URI（base64 内联图片，通常是小图标）
            if src.startswith("data:"):
                continue

            # 相对 URL 转绝对 URL
            img_url = urljoin(base_url, src)

            # SSRF 校验（复用 url_validator，与主 URL 同等防护）
            try:
                from app.core.url_validator import validate_url, URLValidationError
                validate_url(img_url, allow_private=settings.is_development)
            except URLValidationError:
                logger.warning(f"图片 URL 安全校验失败，跳过: {img_url[:100]}")
                continue
            except Exception as e:
                logger.warning(f"图片 URL 校验异常，跳过: {e}")
                continue

            # 下载图片并过滤小图片
            image_bytes = self._download_and_filter_image(img_url)
            if image_bytes is None:
                continue

            # 生成图片文本（OCR + 多模态描述）
            text = self._generate_image_text(image_bytes, img_url, idx, image_processor)
            if text:
                image_texts.append(text)

        return "\n\n".join(image_texts)

    def _download_and_filter_image(self, img_url: str) -> Optional[bytes]:
        """
        下载图片并过滤小图片

        作用：
            下载图片二进制，检测尺寸，小于阈值的跳过（装饰性小图标无检索价值）。
            下载失败或超时时返回 None，不影响其他图片。

        安全防护：
            - 禁用重定向（同主 URL）
            - 大小限制（URL_IMAGE_MAX_SIZE）
            - 尺寸过滤（URL_IMAGE_MIN_SIZE）

        参数：
            img_url: str - 图片 URL（已经过 SSRF 校验）

        返回:
            Optional[bytes] - 图片二进制（小图片或失败时返回 None）
        """
        from app.core.config import settings

        try:
            response = requests.get(
                img_url,
                headers=self._HEADERS,
                timeout=settings.URL_IMAGE_DOWNLOAD_TIMEOUT,
                allow_redirects=False,  # 同主 URL，禁用重定向
                stream=True,
            )
            if response.is_redirect or response.is_permanent_redirect:
                response.close()
                return None

            response.raise_for_status()

            # 检查 Content-Length
            content_length = response.headers.get("Content-Length")
            if content_length and int(content_length) > settings.URL_IMAGE_MAX_SIZE:
                response.close()
                return None

            # 下载图片
            image_bytes = response.content
            response.close()

            if len(image_bytes) > settings.URL_IMAGE_MAX_SIZE:
                return None

            # 检测图片尺寸，过滤小图片
            if not self._is_image_large_enough(image_bytes):
                logger.debug(f"图片尺寸过小，跳过: {img_url[:100]}")
                return None

            return image_bytes

        except Exception as e:
            logger.debug(f"图片下载失败（{img_url[:100]}）: {e}")
            return None

    @staticmethod
    def _is_image_large_enough(image_bytes: bytes) -> bool:
        """
        检测图片尺寸是否大于阈值

        作用：
            用 PIL 读取图片宽高，小于 URL_IMAGE_MIN_SIZE 则视为小图片（图标/装饰）。
            非 PIL 支持格式（如 SVG 无固定尺寸）保守放行。

        参数：
            image_bytes: bytes - 图片二进制数据

        返回:
            bool - True 表示图片足够大（应保留）
        """
        from app.core.config import settings

        try:
            from PIL import Image
            import io

            img = Image.open(io.BytesIO(image_bytes))
            width, height = img.size
            return width * height >= settings.URL_IMAGE_MIN_SIZE
        except Exception:
            # 非 PIL 支持格式（如 SVG）保守放行
            return True

    def _generate_image_text(
        self,
        image_bytes: bytes,
        img_url: str,
        idx: int,
        image_processor: Optional[object],
    ) -> str:
        """
        生成图片文本描述

        作用：
            将图片二进制保存为临时文件，调用 ImageProcessor 的 OCR/Vision
            生成文本描述。OCR/Vision 不可用时仅返回图片 URL 元信息。

        参数：
            image_bytes: bytes - 图片二进制
            img_url: str - 图片 URL（用于元信息）
            idx: int - 图片序号
            image_processor: Optional[ImageProcessor] - 图片处理器

        返回:
            str - 图片文本描述
        """
        import tempfile
        import os

        tmp_path = None
        try:
            # 保存临时文件（ImageProcessor 的 OCR/Vision 接受文件路径）
            with tempfile.NamedTemporaryFile(suffix=".img", delete=False) as f:
                f.write(image_bytes)
                tmp_path = f.name

            ocr_text = ""
            description = ""

            if image_processor is not None:
                from app.core.config import settings
                if settings.ENABLE_OCR:
                    ocr_text = image_processor._ocr_image(tmp_path)
                if settings.ENABLE_VISION:
                    description = image_processor._describe_image(tmp_path)

            # 拼接图片文本
            parts = [f"[图片{idx}]"]
            if ocr_text:
                parts.append(f"OCR文字: {ocr_text}")
            if description:
                parts.append(f"图片描述: {description}")
            if not ocr_text and not description:
                parts.append(f"来源: {img_url}")
            return "\n".join(parts)

        except Exception as e:
            logger.debug(f"图片文本生成失败（{img_url[:100]}）: {e}")
            return f"[图片{idx}] 来源: {img_url}"
        finally:
            if tmp_path and os.path.exists(tmp_path):
                try:
                    os.remove(tmp_path)
                except Exception:
                    pass

    def parse_to_context(
        self,
        ctx: PipelineContext,
        url: Optional[str] = None,
    ) -> None:
        """
        解析并填充上下文

        作用：
            下载并提取正文，写入 ctx.raw_text 和 ctx.pages。

        参数：
            ctx: PipelineContext - 流水线上下文
            url: Optional[str] - URL（为空则用 ctx.file_path）
        """
        target_url = url or ctx.file_path
        text = self.parse(target_url)
        ctx.raw_text = text
        ctx.pages = [
            ParsedPage(
                page_number=1,
                text=text,
                blocks=[TextBlock(text=text, page_number=1)],
                is_blank=not text.strip(),
            )
        ]


# ============================================
# 解析器工厂
# ============================================

# 全局解析器实例（无状态，可复用）
# 作用：避免每次解析都创建新实例
_markdown_parser = MarkdownParser()
_txt_parser = TxtParser()
_docx_parser = DocxParser()
_url_parser = UrlParser()


def get_parser(file_type: str):
    """
    根据文件类型获取解析器

    作用：
        工厂方法，根据扩展名返回对应的解析器实例。

    参数：
        file_type: str - 文件类型（.md/.txt/.docx/.url）

    返回:
        对应的解析器实例，无匹配则返回 None
    """
    file_type = file_type.lower()
    if file_type in (".md", ".markdown"):
        return _markdown_parser
    elif file_type == ".txt":
        return _txt_parser
    elif file_type == ".docx":
        return _docx_parser
    elif file_type in (".url", ".html", ".htm"):
        return _url_parser
    return None
