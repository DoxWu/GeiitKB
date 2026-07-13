"""
文档管理路由模块（生产版）

作用：
    定义文档管理相关的 API 接口，包括：
    - 上传文档（异步触发 Celery 任务处理）
    - 获取文档列表（分页 + 状态筛选）
    - 获取文档详情（含处理进度、质量分）
    - 删除文档（软删除）
    - 从 URL 导入文档
    - 重新处理文档
    - 查询任务状态（前端轮询进度）

实现方式：
    1. 文件上传使用 UploadFile
    2. 上传后触发 Celery 异步任务（不阻塞 HTTP 请求）
    3. 通过 Depends 注入认证和数据库
    4. 删除使用软删除（is_deleted 标记）
    5. 文件哈希去重（相同文件不重复处理）
"""

import os
import uuid
import logging
from fastapi import APIRouter, Depends, HTTPException, status, UploadFile, File, Form, Query, Path, Request
from fastapi.responses import FileResponse, PlainTextResponse
from sqlalchemy.orm import Session
from sqlalchemy.exc import IntegrityError
from sqlalchemy import func
from typing import Any, Optional

from app.core.database import get_db
from app.core.config import settings
from app.core.url_validator import validate_url, URLValidationError, sanitize_filename, validate_document_title
from app.core.rate_limit import rate_limit
from app.core.redis import RedisManager, RedisKeys
from app.api.deps import get_current_active_user, get_current_regular_user
from app.models.user import User
from app.models.document import Document
from app.schemas.document import (
    DocumentResponse,
    DocumentListResponse,
    TaskStatusResponse,
    DocumentProcessSummary,
    DocumentRenameRequest,
)
from app.services.document_processor import document_processor
from app.services.permission import permission_service, VISIBILITY_PRIVATE, VISIBILITY_PUBLIC
from app.services.audit_service import audit_service

logger = logging.getLogger(__name__)

# 创建路由器
router = APIRouter(prefix="/documents", tags=["文档管理"])


# ============================================
# M-4 修复：超级管理员操作审计日志
# ============================================

def _audit_superuser_action(
    action: str,
    superuser_id: int,
    document_id: int,
    owner_id: int,
    extra: Optional[dict] = None,
) -> None:
    """
    记录超级管理员访问/操作他人文档的审计日志（M-4 修复）

    作用：
        超级管理员有权访问/管理任意文档，但这类跨用户操作应留痕，
        便于安全审计追溯。日志结构化记录操作人、被操作对象、属主、动作。

    参数：
        action: str - 操作动作（如 access/delete/reprocess）
        superuser_id: int - 超级管理员用户 ID
        document_id: int - 被操作的文档 ID
        owner_id: int - 文档原始属主 ID
        extra: Optional[dict] - 附加信息（如 IP、查询参数等）
    """
    import json as _json
    audit_entry = {
        "event": "superuser_cross_user_action",
        "action": action,
        "superuser_id": superuser_id,
        "document_id": document_id,
        "owner_id": owner_id,
        "extra": extra or {},
    }
    # 使用 warning 级别确保审计日志被采集（INFO 级别生产可能被过滤）
    logger.warning(f"[AUDIT] {_json.dumps(audit_entry, ensure_ascii=False)}")


# ============================================
# 上传文档
# ============================================

@router.post(
    "/upload",
    response_model=DocumentResponse,
    status_code=status.HTTP_201_CREATED,
    summary="上传文档",
    # 限流：每小时最多 20 次上传
    # 作用：防止单用户大量上传消耗服务器存储和计算资源
    dependencies=[Depends(rate_limit("upload", per_hour=settings.RATE_LIMIT_UPLOAD_PER_HOUR))],
)
async def upload_document(
    file: UploadFile = File(..., description="要上传的文件"),
    # M-5 修复：title 长度校验，防止超过 DB String(200) 导致 500 错误
    title: Optional[str] = Form(None, max_length=200, description="文档标题（最多 200 字符）"),
    category: str = Form(default="other", max_length=50, description="文档分类"),
    tags: str = Form(default="", max_length=500, description="标签，逗号分隔（最多 500 字符）"),
    visibility: str = Form(
        default="private",
        description="可见性：private 个人文档库（默认）/ public 公共文档库（仅管理员）"
    ),
    db: Session = Depends(get_db),
    # 任务5：上传接口限制为正式用户，guest 用户返回 403
    current_user: User = Depends(get_current_regular_user),
) -> Any:
    """
    上传文档接口（生产版）

    作用：
        接收用户上传的文件，保存到服务器，创建数据库记录，
        并触发 Celery 异步任务执行文档处理流水线。

        上传接口立即返回，文档在后台异步处理。
        前端可通过 task_id 轮询任务状态。

    实现方式：
        1. 验证文件类型（白名单）
        2. 清洗文件名（防路径遍历）
        3. 分块写入文件（防内存 DoS，边写边检查大小）
        4. 计算文件哈希（去重）
        5. 创建数据库记录（status=pending，捕获 IntegrityError）
        6. 触发 Celery 异步任务
        7. 立即返回文档信息（含 task_id）

    请求：
        - multipart/form-data 格式
        - file: 文件
        - title: 标题（可选）
        - category: 分类
        - tags: 标签（逗号分隔）

    响应（201）：
        文档信息，含 task_id 用于查询处理进度

    错误：
        400: 文件类型不支持或文件过大
        409: 文件已存在（哈希去重）
    """
    # 1. 验证文件类型（白名单校验）
    # 作用：防止上传可执行文件等危险类型
    file_type = document_processor.get_file_type(file.filename)
    if not document_processor.is_allowed_file_type(file.filename):
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail={
                "error": {
                    "code": "INVALID_FILE_TYPE",
                    "message": f"不支持的文件类型: {file_type}，支持: {', '.join(settings.ALLOWED_FILE_TYPES)}"
                }
            }
        )

    # M-23 修复：MIME 类型双重校验，防止伪造扩展名攻击（evil.exe → evil.pdf）
    # 作用：扩展名校验可被绕过（客户端只改文件名），MIME 校验增加一层防护
    #   content_type 为空时降级通过（兼容部分客户端不传 content_type 的场景）
    if not document_processor.validate_file_mime_type(file.filename, file.content_type):
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail={
                "error": {
                    "code": "MIME_TYPE_MISMATCH",
                    "message": "文件 MIME 类型与扩展名不匹配，疑似伪造文件类型"
                }
            }
        )

    # L-11 说明：文件内容级恶意代码扫描（如 ClamAV）未集成
    # 当前缓解措施：扩展名白名单 + MIME 校验 + 大小限制 + 分块写入防 OOM
    # 生产环境建议：接入 ClamAV 或云厂商内容安全服务对上传文件做异步扫描，
    #   扫描完成前文档状态保持 pending_scan，扫描通过后才进入处理流水线

    # 2. 清洗文件名（防路径遍历攻击）
    # 作用：剥离路径前缀、移除危险字符、防止 ../../../etc/passwd 等攻击
    # 安全要求：file.filename 来自客户端，完全不可信
    safe_filename = sanitize_filename(file.filename) if file.filename else "unnamed"
    unique_filename = f"{uuid.uuid4().hex}_{safe_filename}"
    file_path = os.path.join(settings.UPLOAD_DIR, unique_filename)

    # 确保上传目录存在
    os.makedirs(settings.UPLOAD_DIR, exist_ok=True)

    # 3. 分块写入文件 + 边写边检查大小（防内存 DoS）
    # 作用：避免 await file.read() 一次性将大文件读入内存导致 OOM
    #       分块写入，每块 1MB，累计大小超限立即中止并清理
    chunk_size = 1024 * 1024  # 1MB
    file_size = 0
    try:
        with open(file_path, "wb") as f:
            while True:
                chunk = await file.read(chunk_size)
                if not chunk:
                    break
                file_size += len(chunk)
                # 超过大小限制，立即中止并清理
                if file_size > settings.MAX_FILE_SIZE:
                    # L-15 修复：移除冗余的 f.close()，with 语句退出时自动关闭
                    #   原实现手动 f.close() 后 with 语句退出时会再次 close（已关闭文件句柄）
                    os.remove(file_path)
                    raise HTTPException(
                        status_code=status.HTTP_400_BAD_REQUEST,
                        detail={
                            "error": {
                                "code": "FILE_TOO_LARGE",
                                "message": f"文件过大，最大支持 {settings.MAX_FILE_SIZE // (1024*1024)}MB"
                            }
                        },
                    )
                f.write(chunk)
    except HTTPException:
        raise
    except Exception as e:
        # 写入异常时清理临时文件
        if os.path.exists(file_path):
            os.remove(file_path)
        logger.error(f"文件上传写入失败: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail={"error": {"code": "FILE_WRITE_ERROR", "message": "文件保存失败"}},
        )

    # 4. 计算文件哈希（用于去重）
    # 作用：相同内容的文件不重复处理
    file_hash = document_processor.compute_file_hash(file_path)

    # 5. 检查是否已存在相同哈希的文档
    # M-2 修复：加 Redis 分布式锁防止 TOCTOU 竞态（两个并发请求同时通过检查）
    # 作用：原实现 check-then-insert 无锁，两个并发上传相同文件都通过检查后都插入
    # 修复：用 file_hash 作为锁 key，确保同一哈希的检查+插入是原子操作
    upload_lock_key = RedisKeys.distributed_lock(f"upload:hash:{file_hash}") if file_hash else None
    upload_lock_token = None
    if upload_lock_key:
        upload_lock_token = RedisManager.acquire_lock(upload_lock_key, ttl=30)
        if upload_lock_token is None:
            # 另一个并发上传相同文件的请求正在处理中
            try:
                os.remove(file_path)
            except Exception:
                pass
            raise HTTPException(
                status_code=status.HTTP_409_CONFLICT,
                detail={"error": {"code": "UPLOAD_IN_PROGRESS", "message": "相同文件正在上传中，请稍后重试"}},
            )

    # M-2 修复：try 块覆盖"检查去重 → 创建记录 → DB 插入"全程，finally 释放锁
    try:
        if file_hash:
            existing = db.query(Document).filter(
                Document.file_hash == file_hash,
                Document.is_deleted == False,
                Document.user_id == current_user.id,
            ).first()
            if existing:
                # 删除刚上传的重复文件
                try:
                    os.remove(file_path)
                except Exception:
                    pass
                raise HTTPException(
                    status_code=status.HTTP_409_CONFLICT,
                    detail={
                        "error": {
                            "code": "FILE_ALREADY_EXISTS",
                            "message": f"文件已存在，文档ID: {existing.id}",
                            "document_id": existing.id
                        }
                    },
                )

        # 6. 创建数据库记录
        doc_title = title if title else safe_filename
        tag_list = [tag.strip() for tag in tags.split(",") if tag.strip()] if tags else []

        # 校验可见性（普通用户请求 public 会被降级为 private）
        effective_visibility = permission_service.validate_visibility(visibility, current_user)

        db_document = Document(
            title=doc_title,
            file_name=safe_filename,
            file_path=file_path,
            file_type=file_type,
            file_size=file_size,
            file_hash=file_hash,
            status="pending",
            processing_step="uploaded",
            processing_progress=0,
            user_id=current_user.id,
            visibility=effective_visibility,
            metadata_={"category": category, "tags": tag_list} if tag_list else None,
        )

        # 捕获 IntegrityError（并发上传同名/同哈希时的竞态条件）
        try:
            db.add(db_document)
            db.commit()
            db.refresh(db_document)
        except IntegrityError:
            db.rollback()
            # 清理已写入的文件
            try:
                os.remove(file_path)
            except Exception:
                pass
            raise HTTPException(
                status_code=status.HTTP_409_CONFLICT,
                detail={"error": {"code": "FILE_ALREADY_EXISTS", "message": "文件已存在"}},
            )
    finally:
        # M-2 修复：无论成功失败，释放上传去重锁
        if upload_lock_key and upload_lock_token:
            RedisManager.release_lock(upload_lock_key, upload_lock_token)

    # 7. 触发 Celery 异步任务
    # 作用：上传接口立即返回，文档在后台异步处理
    task_id = None
    try:
        from app.tasks.document_tasks import process_document_task
        task = process_document_task.delay(db_document.id)
        task_id = task.id
    except Exception as e:
        # Celery 触发失败：文档保持 pending，标记 failed 让用户重试
        # 作用：Celery 不可用时，文档无法处理，标记 failed 提示用户重新触发
        logger.error(f"触发 Celery 任务失败，文档将保持 pending 状态: {e}")
        db_document.status = "failed"
        db_document.error_message = "任务触发失败，请稍后重试"  # 脱敏：不暴露内部错误
        db.commit()
        db.refresh(db_document)
        return db_document

    # Celery 已触发，尝试记录 task_id
    # H-4 修复：commit task_id 失败不标记 failed（任务已在后台运行，会自行更新状态）
    # 作用：原实现把 delay 和 commit 放同一 try，commit 失败时误标 failed，
    #       但 Celery 任务会随后把状态改为 processing/completed，造成状态闪烁。
    #       修复后：delay 成功后单独 commit task_id，失败仅记警告并 rollback，
    #       让任务自行更新文档状态（任务内部会设置 processing/completed）。
    try:
        db_document.task_id = task_id
        db.commit()
        db.refresh(db_document)
        logger.info(
            f"文档上传成功，已触发处理任务：doc_id={db_document.id}, task_id={task_id}"
        )
    except Exception as e:
        # task_id 记录失败，但任务已在后台运行，会自行更新文档状态
        logger.warning(
            f"记录 task_id 失败（doc_id={db_document.id}），任务仍在后台运行: {e}"
        )
        db.rollback()
        db.refresh(db_document)

    return db_document


# ============================================
# 获取文档列表
# ============================================

@router.get(
    "",
    response_model=DocumentListResponse,
    summary="获取文档列表"
)
def list_documents(
    # 分页参数校验：page ≥1，page_size 1~100
    # 作用：防止 page=0 或 page_size=999999 导致的异常查询和性能问题
    page: int = Query(default=1, ge=1, description="页码，从 1 开始"),
    page_size: int = Query(default=10, ge=1, le=100, description="每页数量，1-100"),
    # D4-02 游标分页：可选 cursor 参数，传入时使用游标分页（向后兼容 offset/limit）
    # 作用：大数据量时避免 offset 深翻页性能退化（offset 100000 需扫描 100000 行）
    cursor: Optional[int] = Query(
        None, ge=1, description="游标（上一页最后一条文档ID，传入时使用游标分页）"
    ),
    # D2-01 全文检索：替代 LIKE，使用 PostgreSQL to_tsvector + @@ 操作符
    # 作用：利用 GIN 索引加速搜索，同时保留 ilike 作为短词降级兼容
    search: Optional[str] = Query(
        None, max_length=200, description="搜索关键词（标题全文检索）"
    ),
    # M-6 修复：status 参数枚举校验，防止无效状态值导致空查询
    status: Optional[str] = Query(
        None,
        pattern=r"^(pending|processing|completed|failed|low_quality)$",
        description="状态筛选：pending/processing/completed/failed/low_quality"
    ),
    scope: str = Query(
        default="accessible",
        description="范围：accessible 我可访问的（自己的+公共库，默认）/ mine 我上传的 / public 公共文档库"
    ),
    # 修复 Issue 7：添加 folder_id 查询参数，实现分支文档隔离显示
    # 作用：此前 list_documents 不接受 folder_id，导致前端切换分支时仍返回所有文档，
    #       任意分支都能看见几乎所有文档。修复后按 folder_id 精确过滤。
    #       - folder_id 传入具体值：仅返回该分支下的文档
    #       - folder_id 未传（None）：保持原行为（不按分支过滤）
    folder_id: Optional[int] = Query(
        None, ge=1, description="文档库分支ID（传入时仅返回该分支下的文档）"
    ),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_active_user),
) -> Any:
    """
    获取文档列表（分页 + 权限隔离 + 全文检索）

    作用：
        返回当前用户有权查看的文档，支持分页、状态筛选、全文检索和范围选择。
        自动过滤已软删除的文档。

        【权限隔离】通过 scope 控制可见范围：
        - accessible（默认）：自己上传的 + 公共文档库（即可检索的全部）
        - mine：仅自己上传的文档
        - public：仅公共文档库（visibility=public）

        【D4-02 游标分页】传入 cursor 参数时使用游标分页（WHERE id < cursor），
        未传入时保持 offset/limit 分页（向后兼容）。响应含 next_cursor 字段。

        【D2-01 全文检索】search 参数使用 to_tsvector + @@ 全文检索，
        同时保留 ilike 作为短词降级兼容。

        【Issue 7 分支隔离】folder_id 传入时仅返回该分支下的文档，
        未传入时不按分支过滤（保持向后兼容）。

    查询参数：
        - page: 页码，默认 1（cursor 模式下忽略）
        - page_size: 每页数量，默认 10
        - cursor: 游标（可选，传入时使用游标分页）
        - search: 搜索关键词（标题全文检索）
        - status: 状态筛选（pending/processing/completed/failed/low_quality）
        - scope: 范围（accessible/mine/public）
        - folder_id: 文档库分支ID（可选，传入时仅返回该分支下的文档）

    响应（200）：
        {
            "items": [...],
            "total": 100,
            "page": 1,
            "page_size": 10,
            "next_cursor": null
        }
    """
    # D4-03 验证结论：DocumentResponse schema 不含 user 关联字段，
    # 查询仅访问 Document 表字段，无 relationship 懒加载，不存在 N+1 问题。
    # 若未来 DocumentResponse 新增 user 关联字段，需改用 joinedload(User) 预加载。
    query = db.query(Document).filter(Document.is_deleted == False)

    # 按范围过滤（权限隔离核心）
    # 作用：根据 scope 限定可见文档范围
    if scope == "mine":
        # 仅自己上传的
        query = query.filter(Document.user_id == current_user.id)
    elif scope == "public":
        # 仅公共文档库
        query = query.filter(Document.visibility == VISIBILITY_PUBLIC)
    else:
        # accessible（默认）：自己的 + 公共库
        # 超级管理员可看全部
        if current_user.is_superuser:
            pass  # 不加过滤
        else:
            query = query.filter(
                (Document.user_id == current_user.id)
                | (Document.visibility == VISIBILITY_PUBLIC)
            )

    # 状态筛选
    if status:
        query = query.filter(Document.status == status)

    # 修复 Issue 7：分支文档隔离过滤
    # 作用：folder_id 传入时仅返回该分支下的文档，实现分支隔离显示
    # 注意：folder_id 来自前端当前选中的分支，未选中分支时为 None（显示全部）
    if folder_id is not None:
        query = query.filter(Document.folder_id == folder_id)

    # D2-01 全文检索：search 参数使用 to_tsvector + @@ 操作符
    # 作用：利用 PostgreSQL GIN 索引加速搜索，ilike 作为短词降级兼容
    if search:
        search_term = search.strip()
        if search_term:
            query = query.filter(
                func.to_tsvector("simple", Document.title).match(search_term)
                | Document.title.ilike(f"%{search_term}%")
            )

    # 获取总数
    total = query.count()

    # D4-02 游标分页：传入 cursor 时使用游标分页，否则保持 offset/limit（向后兼容）
    next_cursor = None
    if cursor:
        # 游标分页：WHERE id < cursor ORDER BY id DESC LIMIT page_size
        # 作用：避免 offset 深翻页性能退化
        query = query.filter(Document.id < cursor)
        documents = query.order_by(Document.id.desc()).limit(page_size).all()
        # 若返回满页，说明可能还有更多数据，设置 next_cursor
        if len(documents) == page_size:
            next_cursor = documents[-1].id
    else:
        # 原 offset 逻辑保持（向后兼容）
        offset = (page - 1) * page_size
        documents = query.offset(offset).limit(page_size).all()

    return {
        "items": documents,
        "total": total,
        "page": page,
        "page_size": page_size,
        "next_cursor": next_cursor,
    }


# ============================================
# 获取文档详情
# ============================================

@router.get(
    "/{document_id}",
    response_model=DocumentResponse,
    summary="获取文档详情"
)
def get_document(
    # L-4 修复：路径参数正整数校验，防止 document_id=0 或负数导致异常查询
    document_id: int = Path(..., ge=1, description="文档ID（正整数）"),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_active_user),
) -> Any:
    """
    获取文档详情

    作用：
        返回指定文档的详细信息，包括处理进度、质量分等。

    路径参数：
        - document_id: 文档ID

    响应（200）：
        文档详细信息，含 processing_step、processing_progress、quality_score 等

    错误：
        404: 文档不存在或无权访问
    """
    # 权限校验：可访问 = 自己上传的 / 公共文档库 / 超级管理员
    # 作用：支持查看公共文档，同时隔离他人私有文档
    if not permission_service.can_access_document(
        db, current_user.id, document_id, current_user.is_superuser
    ):
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail={"error": {"code": "DOCUMENT_NOT_FOUND", "message": "文档不存在或无权访问"}}
        )

    document = db.query(Document).filter(
        Document.id == document_id,
        Document.is_deleted == False,
    ).first()

    if not document:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail={"error": {"code": "DOCUMENT_NOT_FOUND", "message": "文档不存在"}}
        )

    # M-4 修复：超级管理员访问他人文档时记录审计日志
    # 作用：超管跨用户访问应留痕，便于安全审计追溯
    if current_user.is_superuser and document.user_id != current_user.id:
        _audit_superuser_action(
            action="access",
            superuser_id=current_user.id,
            document_id=document_id,
            owner_id=document.user_id,
        )

    return document


# ============================================
# 删除文档（软删除）
# ============================================

@router.delete(
    "/{document_id}",
    status_code=status.HTTP_204_NO_CONTENT,
    summary="删除文档"
)
def delete_document(
    request: Request,
    # L-4 修复：路径参数正整数校验，防止 document_id=0 或负数导致异常查询
    document_id: int = Path(..., ge=1, description="文档ID（正整数）"),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_active_user),
) -> None:
    """
    删除文档（软删除）

    作用：
        软删除指定文档：
        1. 标记 is_deleted=True（不真正删除，便于恢复）
        2. 从向量数据库删除对应的分块（停止检索）
        3. 保留文件和数据库记录（30 天后由定时任务清理）

    路径参数：
        - document_id: 文档ID

    错误：
        404: 文档不存在或无权访问
        403: 无权管理此文档
    """
    from datetime import datetime

    # 权限校验：管理操作 = 自己上传的 / 超级管理员
    # 作用：公共文档仅上传者和管理员能删除，其他用户只能查看
    if not permission_service.can_manage_document(
        db, current_user.id, document_id, current_user.is_superuser
    ):
        # 不区分"不存在"和"无权"以免泄露文档存在性
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail={"error": {"code": "DOCUMENT_NOT_FOUND", "message": "文档不存在或无权操作"}}
        )

    # 查询文档
    document = db.query(Document).filter(
        Document.id == document_id,
        Document.is_deleted == False,
    ).first()

    if not document:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail={"error": {"code": "DOCUMENT_NOT_FOUND", "message": "文档不存在"}}
        )

    # M-4 修复：超级管理员删除他人文档时记录审计日志
    if current_user.is_superuser and document.user_id != current_user.id:
        _audit_superuser_action(
            action="delete",
            superuser_id=current_user.id,
            document_id=document_id,
            owner_id=document.user_id,
        )

    # C-4 修复：调整操作顺序，先 commit DB 再删向量
    # 作用：原实现先删向量再 commit，若 commit 失败则向量永久丢失但文档仍 active，
    #       导致数据不一致（文档显示可用但检索不到）。
    # 修复后：先 commit 标记软删除（DB 一致性优先），再删向量（失败可重试，不影响一致性）。
    # 数据一致性原则：DB 是 source of truth，向量可重建，不可逆向恢复。
    document.is_deleted = True
    document.deleted_at = datetime.now()
    document.status = "deleted"
    db.commit()  # 先确保 DB 一致，文档标记为已删除

    # 删除向量数据库中的分块（失败不影响主流程，可后续重试或由定时任务清理）
    # 作用：立即停止该文档参与检索；若删除失败，文档已标记 is_deleted，
    #       检索时会过滤掉，不会影响业务；残留向量可由定时任务 cleanup_orphan_vectors 清理
    try:
        from app.services.vector_store import get_vector_store
        vector_store = get_vector_store()
        vector_store.delete_document_chunks(document_id)
    except Exception as e:
        # 警告但不回滚 DB：文档已软删除，向量残留不影响业务（检索时 is_deleted 过滤）
        logger.warning(
            f"删除向量分块失败（doc_id={document_id}），待定时任务清理: {e}"
        )

    logger.info(f"文档已软删除: doc_id={document_id}")

    # D10-01 审计日志：记录文档删除操作
    # 作用：敏感操作留痕，便于安全审计追溯
    audit_service.log(
        db=db,
        user_id=current_user.id,
        action="document.delete",
        resource_type="document",
        resource_id=document_id,
        detail={"title": document.title, "visibility": document.visibility},
        request=request,
    )


# ============================================
# 重新处理文档
# ============================================

@router.post(
    "/{document_id}/reprocess",
    response_model=DocumentResponse,
    summary="重新处理文档"
)
def reprocess_document(
    # L-4 修复：路径参数正整数校验，防止 document_id=0 或负数导致异常查询
    document_id: int = Path(..., ge=1, description="文档ID（正整数）"),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_active_user),
) -> Any:
    """
    重新处理文档

    作用：
        删除旧的分块数据，重新执行文档处理流水线。
        用于文档处理失败后重试、或调整处理参数后重新处理。

    实现方式：
        1. 验证文档存在
        2. 触发 Celery 重新处理任务
        3. 返回文档信息（含新 task_id）

    路径参数：
        - document_id: 文档ID

    错误：
        404: 文档不存在或无权访问
    """
    # 权限校验：重新处理属于管理操作（自己的 / 超级管理员）
    if not permission_service.can_manage_document(
        db, current_user.id, document_id, current_user.is_superuser
    ):
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail={"error": {"code": "DOCUMENT_NOT_FOUND", "message": "文档不存在或无权操作"}}
        )

    document = db.query(Document).filter(
        Document.id == document_id,
        Document.is_deleted == False,
    ).first()

    if not document:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail={"error": {"code": "DOCUMENT_NOT_FOUND", "message": "文档不存在"}}
        )

    # M-4 修复：超级管理员重新处理他人文档时记录审计日志
    if current_user.is_superuser and document.user_id != current_user.id:
        _audit_superuser_action(
            action="reprocess",
            superuser_id=current_user.id,
            document_id=document_id,
            owner_id=document.user_id,
        )

    # 幂等性检查：正在处理中的文档不允许重复触发
    # 作用：防止并发重复触发导致 Celery 任务堆积、重复计费、状态混乱
    if document.status == "processing":
        raise HTTPException(
            status_code=status.HTTP_409_CONFLICT,
            detail={
                "error": {
                    "code": "DOCUMENT_ALREADY_PROCESSING",
                    "message": "文档正在处理中，请等待当前处理完成后再重试",
                    "task_id": document.task_id,
                }
            },
        )

    # 分布式锁：防止并发重复处理请求（TOCTOU 竞态）
    # 作用：在"检查 status"和"触发任务"之间可能有并发请求通过检查
    #       用 Redis SETNX 实现互斥锁，TTL 10 分钟自动释放（防死锁）
    lock_key = RedisKeys.distributed_lock(f"reprocess:doc:{document_id}")
    # H-2 修复：使用 acquire_lock 获取唯一 token，防误删
    # L-2 修复：TTL 使用配置项 REPROCESS_LOCK_TTL，必须 > Celery 任务最长耗时
    lock_token = RedisManager.acquire_lock(lock_key, ttl=settings.REPROCESS_LOCK_TTL)
    if lock_token is None:
        raise HTTPException(
            status_code=status.HTTP_409_CONFLICT,
            detail={"error": {"code": "REPROCESS_LOCKED", "message": "该文档正在被其他请求处理，请稍后重试"}},
        )

    try:
        # 触发 Celery 重新处理任务
        try:
            from app.tasks.document_tasks import reprocess_document_task
            task = reprocess_document_task.delay(document_id)
            document.task_id = task.id
            # H-3 修复：状态设为 processing 而非 pending，关闭 TOCTOU 竞态窗口
            # 作用：原实现设为 pending，finally 释放锁后到 Celery 任务实际启动前存在窗口，
            #       第二个请求的 status=="processing" 检查会通过（此时仍为 pending），导致并发重复处理。
            #       修复后：持锁期间就设为 processing，锁释放后 status 屏障立即生效，第二个请求被 409 挡住。
            #       Celery 任务内部会再次设置 processing（幂等，无副作用）。
            document.status = "processing"
            document.processing_step = "queued"
            document.processing_progress = 0
            document.error_message = None
            db.commit()
            db.refresh(document)

            logger.info(
                f"文档重新处理任务已触发: doc_id={document_id}, task_id={task.id}"
            )
        except Exception as e:
            logger.error(f"触发重新处理任务失败: {e}")
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail={"error": {"code": "TASK_TRIGGER_FAILED", "message": "任务触发失败，请稍后重试"}}
            )
    finally:
        # 释放分布式锁
        # 作用：无论成功或失败，都释放锁，允许后续重试
        # H-2: 使用 release_lock 比对 token，防止误删他人锁
        if lock_token:
            RedisManager.release_lock(lock_key, lock_token)

    return document


# ============================================
# 移动文档到其他分支
# ============================================

@router.patch(
    "/{document_id}/move",
    response_model=DocumentResponse,
    summary="移动文档到其他分支"
)
def move_document(
    # L-4 修复：路径参数正整数校验
    document_id: int = Path(..., ge=1, description="文档ID（正整数）"),
    # 目标分支ID（null 表示移出分支，归入"未分类"）
    folder_id: Optional[int] = Query(
        None, description="目标分支ID（null 表示移出分支，归入未分类）"
    ),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_active_user),
) -> Any:
    """
    移动文档到其他分支

    作用：
        将文档从一个分支移动到另一个分支，或移出分支（归入未分类）。
        用于文档分类管理，支持跨分支迁移。

    实现方式：
        1. 验证文档存在且用户有管理权限
        2. 验证目标分支存在且属于当前用户
        3. 更新文档的 folder_id
        4. 返回更新后的文档信息

    路径参数：
        - document_id: 文档ID

    查询参数：
        - folder_id: 目标分支ID（null 表示移出分支）

    错误：
        404: 文档不存在或无权操作
        400: 目标分支不存在或不属于当前用户
    """
    # 权限校验：移动属于管理操作（自己的 / 超级管理员）
    if not permission_service.can_manage_document(
        db, current_user.id, document_id, current_user.is_superuser
    ):
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail={"error": {"code": "DOCUMENT_NOT_FOUND", "message": "文档不存在或无权操作"}}
        )

    document = db.query(Document).filter(
        Document.id == document_id,
        Document.is_deleted == False,
    ).first()

    if not document:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail={"error": {"code": "DOCUMENT_NOT_FOUND", "message": "文档不存在"}}
        )

    # M-4 修复：超级管理员移动他人文档时记录审计日志
    if current_user.is_superuser and document.user_id != current_user.id:
        _audit_superuser_action(
            action="move",
            superuser_id=current_user.id,
            document_id=document_id,
            owner_id=document.user_id,
            extra={"target_folder_id": folder_id},
        )

    # 验证目标分支（如果指定了 folder_id）
    if folder_id is not None:
        from app.models.document_folder import DocumentFolder
        target_folder = db.query(DocumentFolder).filter(
            DocumentFolder.id == folder_id,
        ).first()
        if not target_folder:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail={"error": {"code": "FOLDER_NOT_FOUND", "message": "目标分支不存在"}}
            )
        # 分支是用户私有的，只能移动到自己的分支
        # 例外：超级管理员可移动到任意分支
        if target_folder.user_id != current_user.id and not current_user.is_superuser:
            raise HTTPException(
                status_code=status.HTTP_403_FORBIDDEN,
                detail={"error": {"code": "FOLDER_ACCESS_DENIED", "message": "无权使用该分支"}}
            )

    # 更新文档的 folder_id
    document.folder_id = folder_id
    db.commit()
    db.refresh(document)

    logger.info(
        f"文档已移动: doc_id={document_id}, target_folder_id={folder_id}, "
        f"operator={current_user.id}"
    )

    return document


# ============================================
# 获取文档内容（预览）- 任务3
# ============================================

# 文本类文件扩展名集合
# 作用：判断文件是否可直接以文本形式返回预览
_TEXT_PREVIEW_TYPES = {"txt", "md", "csv", "json", "text", "markdown"}

# 文本预览最大字节数（10MB）
# 作用：防止超大文本文件一次性读入内存导致 OOM
_TEXT_PREVIEW_MAX_SIZE = 10 * 1024 * 1024


@router.get(
    "/{document_id}/content",
    summary="获取文档内容（预览）",
)
def get_document_content(
    document_id: int = Path(..., ge=1, description="文档ID（正整数）"),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_active_user),
) -> Any:
    """
    获取文档内容（预览）

    作用：
        返回文档的原始内容，供前端 DocumentPreview 组件预览展示。
        按文件类型返回不同响应：
        - 文本类（txt/md/csv/json）：返回纯文本，前端用 <pre> 渲染
        - PDF：返回文件流（FileResponse），前端用 <iframe> 内联预览
        - DOCX：用 python-docx 提取文本返回（依赖不存在时降级为下载）
        - 其他格式：返回 415 Unsupported Media Type

    实现方式：
        1. 权限校验（can_access_document，与 get_document 一致）
        2. 拼接文件绝对路径（防路径遍历：仅取 basename）
        3. 按文件类型分支返回

    路径参数：
        - document_id: 文档ID

    响应：
        - 200: 文档内容（PlainTextResponse 或 FileResponse）
        - 404: 文档不存在或无权访问
        - 415: 不支持预览的文件类型
        - 500: 读取失败

    错误：
        404: 文档不存在或无权访问
        415: 不支持预览的文件类型
    """
    # 权限校验：可访问 = 自己上传的 / 公共文档库 / 超级管理员
    # 作用：与 get_document 保持一致权限，不泄露文档存在性
    if not permission_service.can_access_document(
        db, current_user.id, document_id, current_user.is_superuser
    ):
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail={"error": {"code": "DOCUMENT_NOT_FOUND", "message": "文档不存在或无权访问"}}
        )

    document = db.query(Document).filter(
        Document.id == document_id,
        Document.is_deleted == False,
    ).first()

    if not document:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail={"error": {"code": "DOCUMENT_NOT_FOUND", "message": "文档不存在"}}
        )

    # 拼接绝对路径（防路径遍历：只用 basename，忽略任何路径前缀）
    # 作用：document.file_path 存储的是相对路径，通过 basename 提取文件名后与 UPLOAD_DIR 拼接
    #       防止恶意构造的 file_path 访问上传目录外的文件
    file_abs_path = os.path.join(settings.UPLOAD_DIR, os.path.basename(document.file_path))
    if not os.path.exists(file_abs_path):
        logger.error(f"文档文件不存在: doc_id={document_id}, path={file_abs_path}")
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail={"error": {"code": "FILE_NOT_FOUND", "message": "文件不存在于服务器"}}
        )

    # 统一文件类型（小写、去前导点）
    file_type = document.file_type.lower().lstrip(".")

    # 文本类：读取并返回文本内容
    if file_type in _TEXT_PREVIEW_TYPES:
        # 大文件保护：超过 10MB 拒绝预览，避免 OOM
        file_size = os.path.getsize(file_abs_path)
        if file_size > _TEXT_PREVIEW_MAX_SIZE:
            raise HTTPException(
                status_code=status.HTTP_413_REQUEST_ENTITY_TOO_LARGE,
                detail={
                    "error": {
                        "code": "FILE_TOO_LARGE_FOR_PREVIEW",
                        "message": f"文件过大（{file_size // 1024 // 1024}MB），不支持在线预览，请下载后查看"
                    }
                },
            )
        try:
            with open(file_abs_path, "r", encoding="utf-8", errors="replace") as f:
                content = f.read()
            return PlainTextResponse(content, media_type="text/plain; charset=utf-8")
        except Exception as e:
            logger.error(f"读取文档内容失败: doc_id={document_id}, error={e}")
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail={"error": {"code": "READ_FAILED", "message": "读取文档内容失败"}},
            )

    # PDF：返回文件流（前端 iframe 内联预览）
    # 作用：FileResponse 默认 Content-Disposition: inline，浏览器内联渲染 PDF
    if file_type == "pdf":
        return FileResponse(
            path=file_abs_path,
            media_type="application/pdf",
            filename=document.file_name,
        )

    # DOCX：用 python-docx 提取文本
    # 作用：DOCX 无法直接在浏览器渲染，提取文本后以纯文本形式预览
    if file_type == "docx":
        try:
            from docx import Document as DocxDocument
            doc = DocxDocument(file_abs_path)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            extracted_text = "\n".join(paragraphs) if paragraphs else "（文档内容为空）"
            return PlainTextResponse(extracted_text, media_type="text/plain; charset=utf-8")
        except ImportError:
            # python-docx 未安装，降级为文件下载
            # 作用：不强制引入依赖，未安装时允许用户下载查看
            logger.info("python-docx 未安装，DOCX 预览降级为下载")
            return FileResponse(path=file_abs_path, filename=document.file_name)
        except Exception as e:
            logger.error(f"DOCX 文本提取失败: doc_id={document_id}, error={e}")
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail={"error": {"code": "DOCX_PARSE_FAILED", "message": "DOCX 文档解析失败"}},
            )

    # 不支持预览的文件类型
    raise HTTPException(
        status_code=status.HTTP_415_UNSUPPORTED_MEDIA_TYPE,
        detail={
            "error": {
                "code": "UNSUPPORTED_PREVIEW_TYPE",
                "message": f"不支持预览 {file_type} 类型文件，请下载后查看"
            }
        },
    )


# ============================================
# 重命名文档 - 任务4
# ============================================

@router.patch(
    "/{document_id}/rename",
    response_model=DocumentResponse,
    summary="重命名文档",
)
def rename_document(
    document_id: int = Path(..., ge=1, description="文档ID（正整数）"),
    body: DocumentRenameRequest = ...,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_active_user),
) -> Any:
    """
    重命名文档

    作用：
        修改文档的 title 字段（显示标题），不影响 file_name（原始文件名）。
        新标题经过安全校验（validate_document_title），防止：
        - 特殊字符注入（控制字符、路径分隔符）
        - 恶意代码（XSS：<script> 等 HTML 标签会被 HTML 实体转义）
        - 路径遍历（../ 序列）
        - 不适当内容

    实现方式：
        1. 权限校验（can_manage_document：自己的 / 超级管理员）
        2. 安全校验新标题（validate_document_title）
        3. 更新 title 字段并提交
        4. 超管操作他人文档时记录审计日志

    路径参数：
        - document_id: 文档ID

    请求体：
        {
            "title": "新文档标题"
        }

    响应（200）：
        更新后的文档信息

    错误：
        400: 标题包含非法字符或过长
        404: 文档不存在或无权操作
    """
    # 权限校验：重命名属于管理操作（自己的 / 超级管理员）
    # 作用：公共文档仅上传者和管理员能重命名
    if not permission_service.can_manage_document(
        db, current_user.id, document_id, current_user.is_superuser
    ):
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail={"error": {"code": "DOCUMENT_NOT_FOUND", "message": "文档不存在或无权操作"}}
        )

    document = db.query(Document).filter(
        Document.id == document_id,
        Document.is_deleted == False,
    ).first()

    if not document:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail={"error": {"code": "DOCUMENT_NOT_FOUND", "message": "文档不存在"}}
        )

    # M-4 修复：超级管理员重命名他人文档时记录审计日志
    if current_user.is_superuser and document.user_id != current_user.id:
        _audit_superuser_action(
            action="rename",
            superuser_id=current_user.id,
            document_id=document_id,
            owner_id=document.user_id,
            extra={"old_title": document.title, "new_title": body.title},
        )

    # 安全校验新标题
    # 作用：validate_document_title 会检测危险字符并 HTML 转义，防止 XSS 和路径遍历
    try:
        safe_title = validate_document_title(body.title)
    except ValueError as e:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail={"error": {"code": "INVALID_TITLE", "message": str(e)}},
        )

    # 检查标题是否实际变化（避免无意义的写操作）
    if document.title == safe_title:
        # 标题未变化，直接返回当前文档（幂等）
        return document

    # 更新标题
    old_title = document.title
    document.title = safe_title
    db.commit()
    db.refresh(document)

    logger.info(
        f"文档重命名: doc_id={document_id}, old='{old_title}', new='{safe_title}', "
        f"operator={current_user.id}"
    )

    return document


# ============================================
# 查询任务状态
# ============================================

@router.get(
    "/{document_id}/task-status",
    response_model=TaskStatusResponse,
    summary="查询文档处理任务状态"
)
def get_document_task_status(
    # L-4 修复：路径参数正整数校验，防止 document_id=0 或负数导致异常查询
    document_id: int = Path(..., ge=1, description="文档ID（正整数）"),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_active_user),
) -> Any:
    """
    查询文档处理任务状态

    作用：
        返回 Celery 任务的状态信息，供前端轮询文档处理进度。
        优先返回 Celery 任务状态，若任务不可查则返回文档状态。

    路径参数：
        - document_id: 文档ID

    响应（200）：
        {
            "task_id": "abc-123",
            "status": "SUCCESS",
            "progress": 100,
            "result": {...},
            "error": null
        }

    错误：
        404: 文档不存在或无权访问
    """
    # 权限校验：能访问即可查看任务状态（自己的 / 公共库 / 超级管理员）
    if not permission_service.can_access_document(
        db, current_user.id, document_id, current_user.is_superuser
    ):
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail={"error": {"code": "DOCUMENT_NOT_FOUND", "message": "文档不存在或无权访问"}}
        )

    document = db.query(Document).filter(
        Document.id == document_id,
        Document.is_deleted == False,
    ).first()

    if not document:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail={"error": {"code": "DOCUMENT_NOT_FOUND", "message": "文档不存在"}}
        )

    # 优先查询 Celery 任务状态
    if document.task_id:
        try:
            from app.core.celery_app import get_task_status
            task_info = get_task_status(document.task_id)
            # 修复 Issue 5：用文档记录中的进度覆盖 Celery 进度
            # 作用：document_tasks.py 的进度回调会实时更新 document.processing_progress，
            #       比 Celery 原生进度更精确（Celery 不记录业务进度）。
            #       但当文档状态已为终态（completed/failed）时，直接用 100/0 避免不一致。
            if document.status == "completed":
                task_info["progress"] = 100
            elif document.status == "failed":
                task_info["progress"] = 0
            else:
                # 处理中：使用文档记录的进度（由进度回调实时更新）
                task_info["progress"] = document.processing_progress
            return task_info
        except Exception as e:
            logger.warning(f"查询 Celery 任务状态失败: {e}")

    # 任务不可查时，从文档记录返回状态
    # 作用：Celery 不可用或任务已过期时，仍能返回基本状态
    status_map = {
        "pending": "PENDING",
        "processing": "STARTED",
        "completed": "SUCCESS",
        "failed": "FAILURE",
        "low_quality": "SUCCESS",
        "deleted": "SUCCESS",
    }

    return {
        "task_id": document.task_id or "",
        "status": status_map.get(document.status, "PENDING"),
        "progress": document.processing_progress,
        "result": {
            "document_id": document.id,
            "status": document.status,
            "chunk_count": document.chunk_count,
            "quality_score": document.quality_score,
        } if document.status in ("completed", "low_quality") else None,
        "error": document.error_message,
    }


# ============================================
# 从 URL 导入文档
# ============================================

@router.post(
    "/import-url",
    response_model=DocumentResponse,
    status_code=status.HTTP_201_CREATED,
    summary="从 URL 导入文档",
    # 限流：每小时最多 20 次 URL 导入
    # 作用：防止滥用 URL 导入功能进行 SSRF 探测或 DDoS 外部站点
    dependencies=[Depends(rate_limit("import_url", per_hour=settings.RATE_LIMIT_UPLOAD_PER_HOUR))],
)
def import_from_url(
    url: str = Form(..., description="网页 URL"),
    title: Optional[str] = Form(None, description="文档标题"),
    category: str = Form(default="web", description="文档分类"),
    tags: str = Form(default="", description="标签，逗号分隔"),
    visibility: str = Form(
        default="private",
        description="可见性：private 个人文档库（默认）/ public 公共文档库（仅管理员）"
    ),
    db: Session = Depends(get_db),
    # 任务5：URL 导入接口限制为正式用户，guest 用户返回 403
    current_user: User = Depends(get_current_regular_user),
) -> Any:
    """
    从 URL 导入文档

    作用：
        下载网页内容，保存为文档，并触发处理流水线。
        内置 SSRF 防护，阻止访问内网地址和危险端口。

    实现方式：
        1. SSRF 校验：协议白名单、IP 黑名单、端口黑名单
        2. 下载网页并提取正文
        3. 保存为 .txt 文件
        4. 创建数据库记录（捕获 IntegrityError）
        5. 触发 Celery 异步任务

    请求：
        - application/x-www-form-urlencoded 格式
        - url: 网页 URL
        - title: 标题（可选）
        - category: 分类
        - tags: 标签

    错误：
        400: URL 无效、SSRF 风险或下载失败
    """
    # 1. SSRF 安全校验（核心安全防护）
    # 作用：阻止通过 URL 导入功能访问内网服务、云元数据、数据库等
    # 安全要求：必须在实际 HTTP 请求前校验，是第一道防线
    try:
        validated_url = validate_url(url, allow_private=settings.is_development)
    except URLValidationError as e:
        # M-21 修复：错误信息脱敏，不回显内网拓扑（如 IP 地址）
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail={"error": {"code": "URL_VALIDATION_FAILED", "message": "URL 不合法或存在安全风险"}},
        )

    # M-1 修复：URL 导入幂等性检查
    # 作用：相同 URL 不重复导入，避免创建重复文档
    # 实现：计算 URL 哈希，存入 file_hash 字段，导入前检查是否已存在
    import hashlib
    url_hash = hashlib.sha256(validated_url.encode("utf-8")).hexdigest()
    existing_url_doc = db.query(Document).filter(
        Document.file_hash == url_hash,
        Document.is_deleted == False,
        Document.user_id == current_user.id,
    ).first()
    if existing_url_doc:
        raise HTTPException(
            status_code=status.HTTP_409_CONFLICT,
            detail={
                "error": {
                    "code": "URL_ALREADY_IMPORTED",
                    "message": f"该 URL 的文档已存在，文档ID: {existing_url_doc.id}",
                    "document_id": existing_url_doc.id,
                }
            },
        )

    # 2. 下载网页并提取正文（含图片文本，任务3增强）
    try:
        text = document_processor.extract_from_url(validated_url)
        if not text.strip():
            raise ValueError("网页内容为空")
    except ValueError as e:
        # 任务3：区分 URL_NOT_HTML（非 HTML 资源）和其他 ValueError
        # 作用：URL 指向图片/PDF 等非网页资源时给出明确提示，而非笼统的"下载失败"
        error_msg = str(e)
        if "URL_NOT_HTML" in error_msg:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail={
                    "error": {
                        "code": "URL_NOT_HTML",
                        "message": "URL 指向的不是网页（可能是图片、PDF 等），请提供 HTML 页面链接"
                    }
                },
            )
        # 其他 ValueError（内容为空、大小超限、重定向等）
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail={"error": {"code": "URL_FETCH_FAILED", "message": "网页内容为空或下载失败，请检查 URL 是否可访问"}}
        )
    except Exception as e:
        # 网络错误、超时等
        logger.error(f"URL 导入失败（url={validated_url[:100]}）: {e}", exc_info=True)
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail={"error": {"code": "URL_FETCH_FAILED", "message": "网页内容下载失败，请检查 URL 是否可访问"}}
        )

    # 3. 保存为 .txt 文件
    unique_filename = f"{uuid.uuid4().hex}_url_import.txt"
    file_path = os.path.join(settings.UPLOAD_DIR, unique_filename)
    os.makedirs(settings.UPLOAD_DIR, exist_ok=True)

    with open(file_path, "w", encoding="utf-8") as f:
        f.write(text)

    # 4. 创建数据库记录
    doc_title = title if title else validated_url[:200]
    tag_list = [tag.strip() for tag in tags.split(",") if tag.strip()] if tags else []

    # 校验可见性（普通用户请求 public 会被降级为 private）
    effective_visibility = permission_service.validate_visibility(visibility, current_user)

    db_document = Document(
        title=doc_title,
        file_name=unique_filename,
        file_path=file_path,
        file_type=".txt",
        file_size=len(text.encode("utf-8")),
        status="pending",
        processing_step="uploaded",
        processing_progress=0,
        user_id=current_user.id,
        visibility=effective_visibility,
        file_hash=url_hash,  # M-1：URL 哈希用于去重
        metadata_={
            "category": category,
            "tags": tag_list,
            "source_url": validated_url,
        },
    )

    # 捕获 IntegrityError（并发导入相同 URL 时的竞态条件）
    try:
        db.add(db_document)
        db.commit()
        db.refresh(db_document)
    except IntegrityError:
        db.rollback()
        try:
            os.remove(file_path)
        except Exception:
            pass
        raise HTTPException(
            status_code=status.HTTP_409_CONFLICT,
            detail={"error": {"code": "DOCUMENT_ALREADY_EXISTS", "message": "该 URL 的文档已存在"}},
        )

    # 5. 触发 Celery 异步任务
    try:
        from app.tasks.document_tasks import process_document_task
        task = process_document_task.delay(db_document.id)
        db_document.task_id = task.id
        db.commit()
        db.refresh(db_document)
    except Exception as e:
        logger.error(f"触发 Celery 任务失败: {e}")
        db_document.status = "failed"
        db_document.error_message = "任务触发失败，请稍后重试"  # 脱敏：不暴露内部错误
        db.commit()
        db.refresh(db_document)

    return db_document


# ============================================
# 获取文档统计信息
# ============================================

@router.get(
    "/stats/overview",
    summary="获取文档统计信息"
)
def get_documents_stats(
    scope: str = Query(
        default="accessible",
        description="范围：accessible（默认）/ mine / public，同文档列表"
    ),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_active_user),
) -> Any:
    """
    获取文档统计信息（权限隔离）

    作用：
        返回当前用户可见范围内的文档统计，用于仪表盘展示。
        scope 语义与文档列表接口一致。

    查询参数：
        - scope: accessible（默认）/ mine / public

    响应（200）：
        {
            "total_documents": 100,
            "completed": 80,
            "processing": 5,
            "failed": 10,
            "low_quality": 5,
            "total_chunks": 5000,
            "total_tokens": 1200000,
            "avg_quality_score": 82.5
        }
    """
    from sqlalchemy import func, or_

    # 构建范围过滤条件（与 list_documents 保持一致）
    # 作用：统一权限范围逻辑，避免统计接口绕过权限隔离
    if scope == "mine":
        scope_filter = Document.user_id == current_user.id
    elif scope == "public":
        scope_filter = Document.visibility == VISIBILITY_PUBLIC
    else:
        # accessible：超级管理员看全部，其他人看自己的 + 公共库
        if current_user.is_superuser:
            scope_filter = None  # 不限
        else:
            scope_filter = or_(
                Document.user_id == current_user.id,
                Document.visibility == VISIBILITY_PUBLIC,
            )

    # 基础查询构造器（统一附加 scope + 未删除过滤）
    def _scoped_query():
        q = db.query(Document).filter(Document.is_deleted == False)
        if scope_filter is not None:
            q = q.filter(scope_filter)
        return q

    # 总文档数
    total_documents = _scoped_query().count()

    # 按状态统计
    status_counts = {}
    status_query = db.query(Document.status, func.count(Document.id)).filter(
        Document.is_deleted == False
    )
    if scope_filter is not None:
        status_query = status_query.filter(scope_filter)
    for stat in status_query.group_by(Document.status).all():
        status_counts[stat[0]] = stat[1]

    # 总块数和 Token 数
    totals_query = db.query(
        func.sum(Document.chunk_count),
        func.sum(Document.total_tokens),
        func.avg(Document.quality_score),
    ).filter(Document.is_deleted == False)
    if scope_filter is not None:
        totals_query = totals_query.filter(scope_filter)
    totals = totals_query.first()

    return {
        "total_documents": total_documents,
        "completed": status_counts.get("completed", 0),
        "processing": status_counts.get("processing", 0),
        "pending": status_counts.get("pending", 0),
        "failed": status_counts.get("failed", 0),
        "low_quality": status_counts.get("low_quality", 0),
        "total_chunks": totals[0] or 0,
        "total_tokens": totals[1] or 0,
        "avg_quality_score": round(totals[2], 2) if totals[2] else 0,
    }
